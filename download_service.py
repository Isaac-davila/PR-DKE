import re
import io
import datetime

try:
    from docx import Document
    DOCX_AVAILABLE = True
except:
    DOCX_AVAILABLE = False

class DownloadService:

    @staticmethod
    def get_date_str():
        return datetime.datetime.now().strftime("%Y-%m-%d")

    @staticmethod
    def sanitize_filename(filename):
        return re.sub(r'[^a-zA-Z0-9_]', '', filename.lower().replace(" ", "_"))

    @staticmethod
    def create_txt(markdown_content):
        return (
            markdown_content
            .replace("# ", "")
            .replace("## ", "")
            .replace("---", "")
        )

    @staticmethod
    def create_docx(markdown_content, title):
        doc = Document()
        doc.add_heading(title, 0)

        for line in markdown_content.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line.replace("# ", ""), level=1)
            elif line.startswith("## "):
                doc.add_heading(line.replace("## ", ""), level=2)
            elif line.startswith("- "):
                doc.add_paragraph(line, style="List Bullet")
            else:
                doc.add_paragraph(line)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    @staticmethod
    def build_filename(date_str, prefix, filename, extension):
        return f"{date_str}_{prefix}_{filename}.{extension}"
    
    @staticmethod
    def get_available_formats():
        formats = ["Markdown (.md)", "Text (.txt)"]
        if DOCX_AVAILABLE:
            formats.append("Word (.docx)")
        return formats
    
    @staticmethod
    def generate_file(markdown_content, format_choice, filename, prefix, title):
        safe_filename = DownloadService.sanitize_filename(filename)
        safe_prefix = DownloadService.sanitize_filename(prefix)
        date_str = DownloadService.get_date_str()

        if format_choice == "Markdown (.md)":
            return (
                markdown_content,
                DownloadService.build_filename(date_str, safe_prefix, safe_filename, "md"),
                "text/markdown"
            )

        elif format_choice == "Text (.txt)":
            return (
                DownloadService.create_txt(markdown_content),
                DownloadService.build_filename(date_str, safe_prefix, safe_filename, "txt"),
                "text/plain"
            )

        elif format_choice == "Word (.docx)":
            return (
                DownloadService.create_docx(markdown_content, title),
                DownloadService.build_filename(date_str, safe_prefix, safe_filename, "docx"),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
    @staticmethod
    def get_timestamp():
        return datetime.datetime.now().strftime("%d.%m.%Y • %H:%M Uhr")