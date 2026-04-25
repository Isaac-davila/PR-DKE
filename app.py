import os
import re

import streamlit as st
import time
import io
import datetime

try:
    from docx import Document
    DOCX_AVAILABLE = True
except:
    DOCX_AVAILABLE = False
from pipeline import run_pipeline
from views.views import render_ui, render_sidebar_history
from ai_service import process_with_ai_action
from ai_service import transcribe_audio
from database import (
    supabase,
    save_to_supabase,
    update_transcription_entry,
    get_transcription_history,
    get_current_user,
    login_with_google,
    logout,
    is_admin,
    has_role
)
from views.admin_view import render_admin_view


def main():
    st.set_page_config(page_title="DKE Audio Agent", layout="wide")

    if "result" not in st.session_state:
        st.session_state.result = None
    if "user" not in st.session_state:
        st.session_state.user = None
    if "view" not in st.session_state:
        st.session_state.view = "main"

    if "current_transcript" not in st.session_state: st.session_state.current_transcript = None
    if "current_filename" not in st.session_state: st.session_state.current_filename = None
    if "current_results" not in st.session_state: st.session_state.current_results = []
    if "current_entry_id" not in st.session_state: st.session_state.current_entry_id = None

    # Token Exchange
    if "code" in st.query_params:
        auth_code = st.query_params["code"]
        try:
            res = supabase.auth.exchange_code_for_session({"auth_code": auth_code})
            if res and res.user:
                st.session_state.user = res.user
        except Exception:
            check = get_current_user()
            if check: st.session_state.user = check
        st.query_params.clear()
        st.rerun()

    user = st.session_state.user or get_current_user()

    if not user:
        st.title("🎙️ DKE Audio Agent")
        st.info("Bitte melde dich an.")
        if st.button("Mit Google anmelden"):
            auth_url = login_with_google()
            st.write(f'<meta http-equiv="refresh" content="0;url={auth_url}">', unsafe_allow_html=True)
            st.stop()
        return

    st.session_state.user = user
    user_is_admin = is_admin(user.id) or user.email == "isaac.davila.mendez@gmail.com"

    with st.sidebar:
        st.subheader("Konto")
        st.write(f"👤 {user.email}")

        # Zeige Admin Konsole nur für Admins
        if user_is_admin:
            if st.button("⚙️ Admin-Konsole"):
                st.session_state.view = "admin"
                st.rerun()

        if st.session_state.view != "main":
            if st.button("🏠 Zurück"):
                st.session_state.view = "main"
                st.rerun()

        if st.button("Abmelden"):
            logout()
            st.session_state.user = None
            st.rerun()

        st.divider()
        # Historie laden (Logik für Admin/Basic ist in database.py)
        history_data = get_transcription_history(user_id=user.id, limit=15)
        selected_old_chat = render_sidebar_history(history_data)

    if st.session_state.view == "admin" and user_is_admin:
        render_admin_view(user)
    else:
        selected_action, uploaded_file, pipeline_mode = render_ui()
        if selected_old_chat:
            st.divider()
            st.subheader(f"Historie: {selected_old_chat['filename']}")
            st.write(selected_old_chat['content'])

            timestamp = datetime.datetime.now().strftime("%d.%m.%Y • %H:%M Uhr")
            markdown_content = f"""# 📄 Historie

**Datei:** {selected_old_chat['filename']}  
**Typ:** Gespeicherter Eintrag  
**Heruntergeladen am:** {timestamp}

---

## 📝 Inhalt
{selected_old_chat['content']}
"""

            st.divider()
            st.subheader("📥 Download")
            formats = ["Markdown (.md)", "Text (.txt)"]
            if DOCX_AVAILABLE:
                formats.append("Word (.docx)")

            format_choice_history = st.selectbox(
                "Format auswählen (Historie)",
                formats,
                key="download_format_history"
            )

            clean_filename = os.path.splitext(selected_old_chat['filename'])[0]
            date_str = datetime.datetime.now().strftime("%Y-%m-%d")
            safe_filename = re.sub(r'[^a-zA-Z0-9_]', '', clean_filename.lower().replace(" ", "_"))
            safe_type = "historie"

            if format_choice_history == "Markdown (.md)":
                file_data = markdown_content
                file_name = f"{date_str}_{safe_type}_{safe_filename}.md"
                mime = "text/markdown"

            elif format_choice_history == "Text (.txt)":
                plain_text = markdown_content.replace("# ", "").replace("## ", "")
                file_data = plain_text
                file_name = f"{date_str}_{safe_type}_{safe_filename}.txt"
                mime = "text/plain"

            elif format_choice_history == "Word (.docx)" and DOCX_AVAILABLE:
                doc = Document()
                doc.add_heading("Historie", 0)

                for line in markdown_content.split("\n"):
                    if line.startswith("# "):
                        doc.add_heading(line.replace("# ", ""), level=1)
                    elif line.startswith("## "):
                        doc.add_heading(line.replace("## ", ""), level=2)
                    elif line.startswith("- "):
                        doc.add_paragraph(line, style="List Bullet")
                    else:
                        doc.add_paragraph(line)

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                file_data = buffer
                file_name = f"{date_str}_{safe_type}_{safe_filename}.docx"
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            st.download_button(
                label=f"⬇️ Als {format_choice_history} herunterladen",
                data=file_data,
                file_name=file_name,
                mime=mime,
                key=f"download_history_{selected_old_chat['id']}"
            )
        elif uploaded_file:
            st.audio(uploaded_file)
            if st.button(f"{selected_action} starten"):
                with st.spinner("KI arbeitet..."):
                    try:
                        st.session_state.current_filename = None
                        st.session_state.current_entry_id = None
                        if pipeline_mode == "Groq (nur Transkription)":
                            transcript = transcribe_audio(uploaded_file)
                        elif pipeline_mode == "Groq + Lokale Diarisierung (CPU)":
                            transcript = run_pipeline(uploaded_file, "Transkribieren", "local")
                        elif pipeline_mode == "AssemblyAI (Transkription + Diarisierung)":
                            transcript = run_pipeline(uploaded_file, "Transkribieren", "assembly")
                        else:
                            transcript = transcribe_audio(uploaded_file)
                        st.session_state.current_transcript = transcript
                        st.session_state.current_filename = uploaded_file.name
                        text_result, ai_tags = process_with_ai_action(transcript, selected_action)
                        result = save_to_supabase(uploaded_file.name, text_result, user.id, transcript, ai_tags)
                        if(result.data and len(result.data) > 0):
                            st.session_state.current_entry_id = result.data[0]["id"]
                        st.session_state.current_results.append({"action": selected_action,"result": text_result})
                        st.success("Erledigt!")
                        st.write(text_result)
                        st.session_state.last_result = text_result
                        st.session_state.last_filename = uploaded_file.name
                        st.session_state.last_action = selected_action
                    except Exception as e:
                        st.error(f"Fehler: {e}")
        if st.session_state.current_transcript:
            st.divider()
            st.subheader("Aktueller Chat")
            if st.session_state.current_filename:
                st.caption(f"Datei: {st.session_state.current_filename}")
            for item in st.session_state.current_results:
                st.write(f"**{item['action']}**")
                st.write(item["result"])
            followup_action = st.selectbox("Weitere Aktion auf aktuelles Transkript anwenden",
                ["Zusammenfassen", "Wichtige Punkte extrahieren"],
                key="current_followup_action" )
            if st.button("Weiter mit aktuellem Chat"):
                with st.spinner("KI arbeitet..."):
                    try:
                        text_result, ai_tags = process_with_ai_action(st.session_state.current_transcript, followup_action)
                        if st.session_state.current_entry_id:
                            update_transcription_entry(st.session_state.current_entry_id, text_result, tag_ids=ai_tags)
                        st.session_state.current_results.append({"action": followup_action, "result": text_result})
                        st.success("Neues Ergebnis erstellt!")
                        st.write(text_result)
                    except Exception as e:
                        st.error(f"Fehler: {e}")

        if not selected_old_chat and "last_result" in st.session_state:
            timestamp = datetime.datetime.now().strftime("%d.%m.%Y • %H:%M Uhr")
            markdown_content = f"""# 📄 KI Ergebnis

**Datei:** {st.session_state.last_filename}  
**Aktion:** {st.session_state.last_action}  
**Heruntergeladen am:** {timestamp}

---

## 📝 Inhalt
{st.session_state.last_result}
"""
            
            st.divider()
            st.subheader("📥 Download")
            formats = ["Markdown (.md)", "Text (.txt)"]
            if DOCX_AVAILABLE:
                formats.append("Word (.docx)")

            format_choice = st.selectbox(
                "Format auswählen",
                formats,
                key="download_format"
            )

            clean_filename = os.path.splitext(st.session_state.last_filename)[0]
            date_str = datetime.datetime.now().strftime("%Y-%m-%d")
            safe_filename = re.sub(r'[^a-zA-Z0-9_]', '', clean_filename.lower().replace(" ", "_"))
            safe_action = re.sub(
                r'[^a-zA-Z0-9_]', 
                '', 
                st.session_state.last_action.lower().replace(" ", "_")
            )

            if format_choice == "Markdown (.md)":
                file_data = markdown_content
                file_name = f"{date_str}_{safe_action}_{safe_filename}.md"
                mime = "text/markdown"

            elif format_choice == "Text (.txt)":
                plain_text = markdown_content.replace("# ", "").replace("## ", "")
                file_data = plain_text
                file_name = f"{date_str}_{safe_action}_{safe_filename}.txt"
                mime = "text/plain"

            elif format_choice == "Word (.docx)" and DOCX_AVAILABLE:
                doc = Document()
                doc.add_heading(f"KI Ergebnis – {st.session_state.last_action}", 0)

                # Markdown grob in Word übernehmen
                for line in markdown_content.split("\n"):
                    if line.startswith("# "):
                        doc.add_heading(line.replace("# ", ""), level=1)
                    elif line.startswith("## "):
                        doc.add_heading(line.replace("## ", ""), level=2)
                    elif line.startswith("- "):
                        doc.add_paragraph(line, style="List Bullet")
                    else:
                        doc.add_paragraph(line)

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                file_data = buffer
                file_name = f"{date_str}_{safe_action}_{safe_filename}.docx"
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            st.download_button(
                label=f"⬇️ Als {format_choice} herunterladen",
                data=file_data,
                file_name=file_name,
                mime=mime,
                key=f"download_result_{clean_filename}"
            )
if __name__ == "__main__":
    main()